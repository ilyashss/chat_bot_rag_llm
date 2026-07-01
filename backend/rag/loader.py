import json
import os

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader
)

from langchain_core.documents import Document
from docx import Document as DocxDocument



def load_document(path):

    ext = os.path.splitext(path)[1].lower()


    # PDF
    if ext == ".pdf":

        loader = PyPDFLoader(path)

        return loader.load()



    # TXT
    elif ext == ".txt":

        loader = TextLoader(
            path,
            encoding="utf-8"
        )

        return loader.load()



    # DOCX
    elif ext == ".docx":

        docx = DocxDocument(path)

        text = "\n".join(
            [
                paragraph.text
                for paragraph in docx.paragraphs
            ]
        )

        return [
            Document(
                page_content=text,
                metadata={
                    "source": path
                }
            )
        ]



    # JSON
    elif ext == ".json":

        with open(
            path,
            "r",
            encoding="utf-8"
        ) as f:

            data = json.load(f)


        text = json.dumps(
            data,
            indent=2
        )


        return [
            Document(
                page_content=text,
                metadata={
                    "source": path
                }
            )
        ]



    else:
        raise ValueError(
            f"Unsupported file type: {ext}"
        )